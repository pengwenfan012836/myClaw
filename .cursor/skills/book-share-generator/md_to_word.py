#!/usr/bin/env python3
"""将Markdown内容转换为Word文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import sys

def create_book_share_word(md_content: str, output_path: str):
    """将markdown内容转换为Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    lines = md_content.strip().split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # 标题处理
        if line.startswith('# '):
            # 一级标题
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.bold = True

        elif line.startswith('## '):
            # 二级标题
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True

        elif line.startswith('### '):
            # 三级标题
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True

        elif line.startswith('!['):
            # 图片（跳过，因为Word中图片路径需要特殊处理）
            pass

        elif line.startswith('|'):
            # 表格处理
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # 解析简单表格
            if len(table_lines) >= 2 and '|' in table_lines[0]:
                rows_data = []
                for tline in table_lines:
                    cells = [c.strip() for c in tline.split('|')[1:-1]]
                    rows_data.append(cells)

                # 判断是否为表头（第一行是分隔符前的行）
                # 简化处理：假设前两行是表头，后面是数据
                if len(rows_data) >= 2:
                    num_cols = len(rows_data[0])
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    table.style = 'Table Grid'

                    for row_idx, row_data in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            if row_idx == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0, 51, 102)
            i -= 1  # 补偿for循环的增量

        elif line.startswith('---'):
            # 分隔线用空段落代替
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith('>'):
            # 引用块
            content = line[1:].strip()
            if content.startswith('#'):
                # 引用中的标题
                content = content.replace('#', '')
                p = doc.add_paragraph(content, style='Quote')
            else:
                p = doc.add_paragraph(content, style='Quote')
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)

        elif line.startswith('**') and line.endswith('**'):
            # 粗体文本
            content = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.bold = True

        elif line.startswith('- '):
            # 无序列表
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])

        elif line[0].isdigit() and '.' in line[:3]:
            # 有序列表
            content = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(content)

        else:
            # 普通文本
            # 处理行内格式
            processed = process_inline_formats(line)
            p = doc.add_paragraph()
            for text, bold, italic in processed:
                run = p.add_run(text)
                run.font.bold = bold
                run.font.italic = italic

        i += 1

    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")

def process_inline_formats(text: str):
    """处理行内格式，返回(文本, 粗体, 斜体)的列表"""
    result = []
    # 简单处理**bold**和*italic*
    remaining = text
    while remaining:
        # 查找加粗
        bold_match = re.search(r'\*\*(.+?)\*\*', remaining)
        italic_match = re.search(r'\*(.+?)\*', remaining)

        if bold_match and (not italic_match or bold_match.start() < italic_match.start()):
            if bold_match.start() > 0:
                result.append((remaining[:bold_match.start()], False, False))
            result.append((bold_match.group(1), True, False))
            remaining = remaining[bold_match.end():]
        elif italic_match:
            if italic_match.start() > 0:
                result.append((remaining[:italic_match.start()], False, False))
            result.append((italic_match.group(1), False, True))
            remaining = remaining[italic_match.end():]
        else:
            result.append((remaining, False, False))
            break

    return result if result else [(text, False, False)]

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python md_to_word.py <输入.md> <输出.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    create_book_share_word(content, output_file)
